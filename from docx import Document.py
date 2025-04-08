from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_header_table(doc):
    """Create the header table with prepared for/prepared by information"""
    table = doc.add_table(rows=2, cols=2)
    
    # Set column widths
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)
    
    # Left cell (Prepared for)
    left_cell = table.cell(0, 0)
    left_cell.text = "Prepared for: Dean's Council & Other Stakeholders\nTitle: Keyano College Enrolment KPI Update\nDate prepared: October 02, 2024"
    
    # Right cell (Prepared by)
    right_cell = table.cell(0, 1)
    right_cell.text = "Prepared by: Bill Guo\nTitle: Institutional Research Analyst\nDepartment: Institutional Research"
    
    # Merge bottom row for subject
    bottom_cell = table.cell(1, 0)
    bottom_cell.merge(table.cell(1, 1))
    bottom_cell.text = "SUBJECT: KPI Updates for Applications and Enrolments as of October 1st, 2024."
    
    # Format header text
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

def create_applications_table(doc):
    """Create the applications comparison table"""
    doc.add_heading('Applications Updates', level=2)
    
    # Create table with 4 rows (header + 3 data) and 10 columns
    table = doc.add_table(rows=4, cols=10)
    table.style = 'Table Grid'
    
    # Set column widths (approximate based on template)
    col_widths = [Inches(1.2)] + [Inches(0.7) for _ in range(9)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Applications\nComparison - Point\nin Time'
    hdr_cells[1].text = 'Summer\n2024'
    hdr_cells[2].text = 'Summer\n2023'
    hdr_cells[3].text = '%\nChange'
    hdr_cells[4].text = 'Fall\n2024'
    hdr_cells[5].text = 'Fall\n2023'
    hdr_cells[6].text = '%\nChange'
    hdr_cells[7].text = 'Winter\n2025'
    hdr_cells[8].text = 'Winter\n2024'
    hdr_cells[9].text = '%\nChange'
    
    # Data rows
    data = [
        ['Population', 'Total Applications', '1433', '39', '3574.4%', '5332', '3731', '42.9%', '1277', '4079', '-68.7%'],
        ['', 'International', '1395', '1', '139400.0%', '3739', '2183', '71.3%', '1174', '3944', '-70.2%'],
        ['', 'Domestic', '38', '38', '0.0%', '1593', '1548', '2.9%', '103', '135', '-23.7%']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_cells):  # Ensure we don't exceed cell count
                row_cells[j].text = str(cell_data)
    
    # Center align all cells except first column
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i > 0:  # Skip first column
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_application_details(doc):
    """Add the detailed application information sections"""
    # Summer 2024 Applications
    doc.add_heading('Summer 2024 Applications', level=3)
    summer_points = [
        "As of October 1, 2024, there are 1426 unique applicants representing 1433 total applications for Summer 2024. The 1426 unique applicants in Summer 2024 are 3556.4% higher compared to the same point in time for Summer 2023 unique applicants.",
        "Of the 1433 total applications for Summer 2024, 352 have accepted offers and paid their admission deposit, representing 24.6%.",
        "The number of international applications is 1395, representing 97.3% of total applications."
    ]
    for point in summer_points:
        p = doc.add_paragraph(point, style='List Bullet')
    
    # Fall 2024 Applications
    doc.add_heading('Fall 2024 Applications', level=3)
    fall_points = [
        "As of October 1, 2024, there are 4914 unique applicants representing 5332 total applications for Fall 2024. The 4914 unique applicants in Fall 2024 are 49.4% higher compared to the same point in time for Fall 2023 unique applicants.",
        "Of the 5332 total applications for Fall 2024, 1261 have accepted offers and paid their admission deposit, representing 23.6%.",
        "The number of international applications is 3739, representing 70.1% of total applications."
    ]
    for point in fall_points:
        p = doc.add_paragraph(point, style='List Bullet')
    
    # Winter 2025 Applications
    doc.add_heading('Winter 2025 Applications', level=3)
    winter_points = [
        "As of October 1, 2024, there are 1252 unique applicants representing 1277 total applications for Winter 2025. The 1252 unique applicants in Winter 2025 are 68.9% lower compared to the same point in time for Winter 2024 unique applicants.",
        "Of the 1277 total applications for Winter 2025, 515 have accepted offers and paid their admission deposit, representing 40.3%.",
        "The number of international applications is 1174, representing 91.9% of total applications."
    ]
    for point in winter_points:
        p = doc.add_paragraph(point, style='List Bullet')

def create_enrolment_table(doc):
    """Create the enrolment comparison table"""
    doc.add_heading('2024-25 Enrolment Updates', level=2)
    
    # Create table with 12 rows (2 headers + 10 data)
    table = doc.add_table(rows=12, cols=8)
    table.style = 'Table Grid'
    
    # Set column widths
    col_widths = [Inches(1.2), Inches(1.2), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Header row (merged)
    hdr_cell = table.rows[0].cells[0]
    hdr_cell.merge(table.rows[0].cells[7])
    hdr_cell.text = "2024-25 Year Start Enrolment Comparison - Point in Time"
    hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Column headers
    col_headers = table.rows[1].cells
    col_headers[0].text = "Category"
    col_headers[1].text = "2024-10-01\n(2024-25)"
    col_headers[2].text = "2023-10-02\n(2023-24)"
    col_headers[3].text = "% Change\nin UHC"
    col_headers[4].text = "% Change\nin FLE"
    col_headers[5].text = "UHC"
    col_headers[6].text = "FLE"
    col_headers[7].text = "UHC\nFLE"
    
    # Data rows
    data = [
        ["By Demographic", "Total Domestic & International", "3232", "2009.075", "2532", "1389.977", "27.6%", "44.5%"],
        ["", "International", "1609", "1156.609", "992", "605.125", "62.2%", "91.1%"],
        ["", "Domestic", "1623", "852.466", "1540", "784.852", "5.4%", "8.6%"],
        ["", "Indigenous", "94", "37.786", "84", "37.631", "11.9%", "0.4%"],
        ["", "Apprenticeship", "433", "118.353", "382", "106.388", "13.4%", "11.2%"],
        ["By Credential", "Certificate", "925", "394.391", "820", "351.228", "12.8%", "12.3%"],
        ["", "Diploma", "1667", "1227.763", "1060", "663.878", "57.3%", "84.9%"],
        ["", "Non-Credential", "675", "386.921", "665", "374.871", "1.5%", "3.2%"],
        ["By Term", "Fall", "2795", "1236.307", "1971", "799.593", "41.8%", "54.6%"],
        ["", "Winter", "1298", "549.372", "1343", "519.388", "-3.4%", "5.8%"]
    ]
    
    for i, row_data in enumerate(data, 2):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    
    # Center align all cells except first column
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i > 0:  # Skip first column
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_enrolment_projections(doc):
    """Add the enrolment projections section"""
    doc.add_heading('Enrolment Actuals vs Projections', level=3)
    doc.add_paragraph('(this section excludes Power Engineering CML, LINC, and Apprenticeship)')
    
    points = [
        "The projected FLE for the 2024-25 academic year is 2475.587. As of October 1, 2024, the actual FLE is 2009.075; this indicates an 81.2% of the projection achieved for the whole academic year.",
        "Fall 2024 (Part-time) *: The projected part-time headcount for Fall 2024 is 99. As of October 1, 2024, the actual headcount is 107, this indicates an 8.1% surpass of the projection achieved for the semester.",
        "Fall 2024 (Full-time) *: The projected full-time headcount for Fall 2024 is 1959. As of October 1, 2024, the actual headcount is 2202; this indicates a 12.4% surpass of the projection achieved for the semester.",
        "Winter 2025 (Part-time) *: The projected part-time headcount for Winter 2025 is 137. As of October 1, 2024, the actual headcount is 104; this indicates 75.9% of the projection achieved for the semester.",
        "Winter 2025 (Full-time) *: The projected full-time headcount for Winter 2025 is 2157. As of October 1, 2024, the actual headcount is 1029, this indicates a 47.7% of the projection achieved for the semester.",
        "It is important to note the following:",
        "i) Enrolment is ongoing for the 2024-25 academic year.",
        "ii) LINC registration takes place in early September with twice-a-month registration until the last intake in May 2025."
    ]
    
    for point in points[:5]:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph(points[5])
    for point in points[6:]:
        doc.add_paragraph(point, style='List Bullet 2')

def add_headcount_projections(doc):
    """Add the headcount projections table"""
    doc.add_heading('2024-25 Enrolments Projections Progress by Unique Headcount', level=2)
    
    # Create table with 7 rows (3 headers + 4 data)
    table = doc.add_table(rows=7, cols=12)
    table.style = 'Table Grid'
    
    # Set column widths
    col_widths = [Inches(1), Inches(1)] + [Inches(0.7) for _ in range(10)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    
    # Header row (merged)
    hdr_cell = table.rows[0].cells[0]
    hdr_cell.merge(table.rows[0].cells[11])
    hdr_cell.text = "2024-25 Enrolments Projections Progress by Unique Headcount"
    hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subheader row
    sub_hdr = table.rows[1].cells
    sub_hdr[0].text = "Population"
    sub_hdr[1].text = "Enrolment Status"
    
    # Merge term headers
    for i, term in enumerate(["Summer", "Fall", "Winter"]):
        start_col = 2 + i*3
        term_cell = table.rows[1].cells[start_col]
        term_cell.merge(table.rows[1].cells[start_col+2])
        term_cell.text = term
        term_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Column headers
    col_headers = table.rows[2].cells
    col_headers[0].text = ""
    col_headers[1].text = ""
    
    for i in range(3):
        start_col = 2 + i*3
        col_headers[start_col].text = "Projected"
        col_headers[start_col+1].text = "Actual"
        col_headers[start_col+2].text = "% Projections Achieved"
    
    # Data rows
    data = [
        ["Domestic", "Full time", "80", "102", "127.5%", "704", "729", "103.6%", "760", "502", "66.1%"],
        ["", "Part time", "4", "52", "1300.0%", "96", "80", "83.3%", "88", "46", "52.3%"],
        ["International", "Full time", "212", "278", "131.1%", "1286", "1473", "114.5%", "1399", "527", "37.7%"],
        ["", "Part time", "0", "9", "N/A", "3", "27", "900.0%", "45", "58", "128.9%"]
    ]
    
    for i, row_data in enumerate(data, 3):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    
    # Center align all cells except first two columns
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i > 1:  # Skip first two columns
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_apprenticeship_section(doc):
    """Add the apprenticeship section"""
    doc.add_heading('Apprenticeship', level=3)
    
    points = [
        "The projected unique headcount for Fall 2024 is 250. As of October 1, 2024, the actual unique headcount is 234; this indicates a 93.6% of the projection achieved for the semester. The following Programs have achieved or surpassed their projection:",
        "- Electrician - Second year, and Third year",
        "- Industrial Mechanic (Millwright) - First Year",
        "- Steamfitter - Pipefitter-Second year",
        "As for Winter 2025, the projected unique headcount is 228, and as of October 1, 2024, the actual unique headcount is 170; this indicates a 74.6% of the projection achieved for the semester. The following Programs have achieved or surpassed their projection:",
        "- Electrician - First year, and Fourth year",
        "- Industrial Mechanic (Millwright) - Third year",
        "- Welder -- First Year"
    ]
    
    doc.add_paragraph(points[0])
    for point in points[1:4]:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph(points[4])
    for point in points[5:]:
        doc.add_paragraph(point, style='List Bullet')

def add_international_section(doc):
    """Add the international students section"""
    doc.add_heading('International Students', level=3)
    
    points = [
        "As of October 1, 2024, there are 1609 unique international students with FLE 1156.609 with high enrolment numbers in Business Administration Diploma -- Management. Increased numbers are visible in Business Administration Diploma - Accounting, Business Administration Diploma - Management Co-op, Business Administration Diploma - Human Resources Management, and Early Learning and Child Care Diploma.",
        "Current International FLE represents 57.6% of the actual Keyano College total FLE."
    ]
    
    for point in points:
        doc.add_paragraph(point)

def add_indigenous_section(doc):
    """Add the indigenous students section"""
    doc.add_heading('Indigenous Students', level=3)
    
    points = [
        "As of October 1, 2024, there are 94 unique Indigenous students with FLE 37.786 with high enrolment in Apprenticeship - Heavy Equipment Technician.",
        "Current Indigenous FLE represents 1.9% of the actual Keyano College total FLE."
    ]
    
    for point in points:
        doc.add_paragraph(point)

def create_briefing_note():
    """Create the complete briefing note document"""
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add header information
    create_header_table(doc)
    
    # Add background section
    doc.add_heading('BACKGROUND:', level=2)
    doc.add_paragraph('The KPI update has been prepared to aid and offer insights to stakeholders regarding the progress of applications and enrolment throughout the academic year. For more detailed information, please refer to the attached KPI Update workbooks.')
    
    # Add current status section
    doc.add_heading('CURRENT STATUS:', level=2)
    
    # Add applications data
    create_applications_table(doc)
    add_application_details(doc)
    
    # Add enrolment data
    create_enrolment_table(doc)
    add_enrolment_projections(doc)
    add_headcount_projections(doc)
    
    # Add additional sections
    add_apprenticeship_section(doc)
    add_international_section(doc)
    add_indigenous_section(doc)
    
    # Save document
    doc.save('Briefing_Note_KC_Applications_and_Enrolments_KPI_Update.docx')

if __name__ == '__main__':
    create_briefing_note()